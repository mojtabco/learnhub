from django.conf import settings
from django.contrib.auth.mixins import LoginRequiredMixin
from django.db.models import Count, ProtectedError,Exists, OuterRef,Q
from django.http import JsonResponse, HttpResponse, Http404
from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse
from django.views import View
from courses.forms import CreateCourseForm, CourseSeasonForm, LessonForm, SearchForm, CourseAccessForm
from courses.models.course_models import CourseAccess, Quote
from courses.models import Course, CourseSeason, Lesson, SubCategory, Category
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from courses.models.lesson_models import StudentProgress, Certificate
from courses.views import number_to_words, paginate_queryset, is_demo_user, restrict_demo_user
from docx import Document  #python-docx
from docx2pdf import convert
import os ,random
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework import serializers


class SearchView(View):
    template_name = "courses/search.html"

    def get_context_data(self, form_search, course_queryset, category, link):
        breadcrumbs = [
            {"label": "دسته بندی", "url": "/"},
            {"label": "داشبورد", "url": "/"},
        ]
        if category:
            breadcrumbs.append({"label": category.title, "url": f"/category/{category.id}"})
        return {'form_search': form_search,'list_course': course_queryset.distinct(),'category': category,
                                                                                    'link': link,'breadcrumbs': breadcrumbs
                                                                                    }

    def get(self, request, *args, **kwargs):
        category = None
        form_search = SearchForm(request.GET)
        search_text = form_search.cleaned_data.get("SearchText") if form_search.is_valid() else None
        link = request.GET.get('link', '')

        course_queryset = self.get_course_queryset(request)

        if link == "owner":
            if request.user.is_authenticated:
                course_queryset = course_queryset.filter(owner=request.user).order_by('-updated')
            else:
                return redirect('login')
        elif link == "tags":
            if request.user.is_authenticated:
                course_queryset = course_queryset.filter(tag=request.user).order_by('-updated')
            else:
                return redirect('login')
        elif link == "like":
            course_queryset = course_queryset.annotate(like_count=Count('likes')).filter(like_count__gt=0).order_by('-like_count', '-updated')
        elif link:
            subcategory_id = SubCategory.objects.filter(title=link).values_list('id', flat=True).first()
            if not subcategory_id:
                raise Http404("دسته بندی یافت نشد")
            course_queryset = course_queryset.filter(subcategory=subcategory_id)
            category = Category.objects.filter(fkey_subcategory=subcategory_id).first()

        if search_text:
            course_queryset = course_queryset.filter(title__icontains=search_text)

        context = self.get_context_data(form_search, course_queryset, category, link)
        return render(request, self.template_name, context)

    def get_course_queryset(self, request):
        if request.user.is_authenticated:
            course_access_queryset = CourseAccess.objects.filter(course=OuterRef('pk')).filter(
                Q(access_users=request.user) | Q(access_groups__in=request.user.profileuser.access_groups.all())
            )
        else:
            course_access_queryset = CourseAccess.objects.filter(course=OuterRef('pk')).filter(access_groups__permission=settings.COURSE_EVERYONE)

        return Course.objects.filter(lesson_count__gt=0, is_published=True).annotate(has_access=Exists(course_access_queryset)).filter(has_access=True)

class ViewView(LoginRequiredMixin,View):
    login_url = 'login'
    template_name = "courses/view.html"

    def get_context_data(self, course, seasons_data, lessons_data, video_season, video_lesson,
                         completed_lessons, total_lessons, course_progress_percentage, user_has_all_access,
                         user_has_download_access):
        breadcrumbs = [
            {"label": "نمایش", "url": ""},
            {"label": "دسته بندی", "url": "/courses/search"},
            {"label": "داشبورد", "url": "/"},
        ]
        return {
            'course': course,
            'seasons_data': seasons_data,
            'lessons_data': lessons_data,
            'video_season': video_season,
            'video_lesson': video_lesson,
            'completed_lessons': completed_lessons,
            'course_progress_percentage': course_progress_percentage,
            'user_has_all_access': user_has_all_access,
            'user_has_download_access': user_has_download_access,
            'breadcrumbs': breadcrumbs
        }

    def get(self, request, cid=None, lid=None, *args, **kwargs):
        try:
            course = Course.objects.get(pk=cid)

            progress, created = StudentProgress.objects.get_or_create(student=request.user, course=course)
            completed_lessons = progress.count_completed_lessons()
            total_lessons = Lesson.objects.filter(course_season__course=course).count()

        except Course.DoesNotExist:
            raise Http404("فصل آموزشی انتخابی در دسترس نیست")

        seasons_data = course.fkey_course_season.select_related().prefetch_related('fkey_lessons').order_by('order')
        lessons_data = []
        for season in seasons_data:
            lessons_data.extend(season.fkey_lessons.all().order_by('order'))

        if lid:
            try:
                video_lesson = Lesson.objects.get(pk=lid)
                video_season = CourseSeason.objects.get(fkey_lessons=video_lesson)
            except Lesson.DoesNotExist:
                raise Http404("آموزشی انتخابی در دسترس نیست")
            except CourseSeason.DoesNotExist:
                raise Http404("فصل آموزشی انتخابی در دسترس نیست")
        else:
            video_season = seasons_data.filter(fkey_lessons__isnull=False).first()
            if video_season:
                video_lesson = video_season.fkey_lessons.first()
            else:
                video_lesson = None

        if not video_lesson:
            raise Http404("هیچ ویدیویی برای نمایش یافت نشد")

        user_has_all_access = request.user.profileuser.has_permission('All-Access')
        user_has_download_access = request.user.profileuser.has_permission('Download')
        course_progress_percentage = (completed_lessons / total_lessons * 100) if total_lessons > 0 else 0

        context = self.get_context_data(
            course, seasons_data, lessons_data, video_season, video_lesson,
            completed_lessons, total_lessons, course_progress_percentage, user_has_all_access, user_has_download_access
        )

        return render(request, self.template_name, context)

class MyTagView(LoginRequiredMixin, View):
    login_url = 'login'
    template_name = "courses/user-page/mytag.html"

    def get_context_data(self,request, courses_data):

        # Pagination
        courses_data = paginate_queryset(courses_data, request, per_page=6)

        breadcrumbs = [
            {"label": "انتخاب های من", "url": ""},
            {"label": "داشبورد", "url": "/"},
        ]
        return {
            'courses_data': courses_data,
            'breadcrumbs': breadcrumbs
        }

    def get(self, request, *args, **kwargs):
        try:
            courses_data = Course.objects.filter(tag=request.user).all().order_by('-updated')
        except Course.DoesNotExist:
            raise Http404("مشکلی در دورهای انتخابی شما پیش امده است")

        context = self.get_context_data(request,courses_data)
        return render(request, self.template_name, context)

class MyCourseView(LoginRequiredMixin, View):
    login_url = 'login'
    template_name = "courses/user-page/mycourse.html"

    def get_search_text(self, request):
        form_search = SearchForm(request.GET)
        search_text = form_search.cleaned_data.get("SearchText") if form_search.is_valid() else None
        return form_search, search_text

    def get_context_data(self, request, search_text=None):
        form = CreateCourseForm()
        user_has_all_access = request.user.profileuser.has_permission('All-Access')
        user_has_upload_access = request.user.profileuser.has_permission('Upload')

        # فیلتر کردن دوره‌ها
        course_queryset = Course.objects.filter(owner=request.user).order_by('-updated')
        if search_text:
            course_queryset = course_queryset.filter(
                Q(title__icontains=search_text) |
                Q(description__icontains=search_text)
            )

        # Pagination
        courses_data = paginate_queryset(course_queryset, request, per_page=6)

        breadcrumbs = [
            {"label": "فایل های من", "url": ""},
            {"label": "داشبورد", "url": "/"},
        ]

        return {
            'form': form,
            'form_search': SearchForm(request.GET),
            'courses_data': courses_data,
            'user_has_all_access': user_has_all_access,
            'user_has_upload_access': user_has_upload_access,
            'breadcrumbs': breadcrumbs
        }

    def get(self, request, *args, **kwargs):
        form_search, search_text = self.get_search_text(request)
        context = self.get_context_data(request, search_text)
        return render(request, self.template_name, context)

    def post(self, request, *args, **kwargs):
        form = CreateCourseForm(request.POST, request.FILES)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.owner = request.user
            instance.save()
        else:
            messages.error(request, 'لطفاً فرم را به درستی پر کنید.')

        return redirect('mycourse')

class MySeasonView(LoginRequiredMixin, View):
    login_url = 'login'
    template_name = "courses/user-page/myseason.html"
    myseason_url = 'myseason'

    def setup(self, request, *args, **kwargs):
        super().setup(request, *args, **kwargs)
        try:
            self.course = Course.objects.get(pk=kwargs['id'])
        except Course.DoesNotExist:
            raise Http404("دوره انتخابی وجود ندارد")

    def get_context_data(self):
        next_order = self.course.fkey_course_season.count() + 1
        # next_order = CourseSeason.objects.filter(course=course).count() + 1
        form = CourseSeasonForm(initial={'order': next_order, 'title': "فصل " + number_to_words(next_order)})
        seasons_data = self.course.fkey_course_season.all().order_by('order')

        breadcrumbs = [
            {"label": "فصل ها", "url": ""},
            {"label": "فایل های من", "url": "/courses/"},
            {"label": "داشبورد", "url": "/"},
        ]

        return {
            'form': form,
            'course': self.course,
            'seasons_data': seasons_data,
            'user_has_all_access': self.course.owner.profileuser.has_permission('All-Access'),
            'user_has_upload_access': self.course.owner.profileuser.has_permission('Upload'),
            'breadcrumbs': breadcrumbs
        }

    def get(self, request, *args, **kwargs):
        context = self.get_context_data()
        return render(request, self.template_name, context)

    def post(self, request, *args, **kwargs):

        form = CourseSeasonForm(request.POST, request.FILES)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.course = self.course
            instance.order = self.course.fkey_course_season.count() + 1 #CourseSeason.objects.filter(course=id).count() + 1
            instance.save()
            return redirect(reverse(self.myseason_url, kwargs={'id': self.course.id}))
        else:
            messages.error(request, 'لطفاً فرم را به درستی پر کنید.')

        context = self.get_context_data()
        return render(request, self.template_name, context)

class MyLessonView(LoginRequiredMixin, View):
    login_url = 'login'
    template_name = "courses/user-page/mylesson.html"
    mylesson_url = 'mylesson'

    def setup(self, request, *args, **kwargs):
        super().setup(request, *args, **kwargs)
        try:
            self.season = CourseSeason.objects.get(pk=kwargs['id'])
            self.course = self.season.course
        except CourseSeason.DoesNotExist:
            raise Http404("فصل انتخابی وجود ندارد")
        except AttributeError:
            raise Http404("دوره مربوط به فصل وجود ندارد")

    def get_context_data(self, request):
        user_has_all_access = request.user.profileuser.has_permission('All-Access')
        user_has_upload_access = request.user.profileuser.has_permission('Upload')

        # next_order = Lesson.objects.filter(course_season=season).count() + 1
        next_order = self.season.fkey_lessons.count() + 1
        form = LessonForm(initial={'order': next_order})
        lessons_data = self.season.fkey_lessons.all().order_by('order')

        # Pagination
        lessons_data = paginate_queryset(lessons_data, request, per_page=6)

        breadcrumbs = [
            {"label": "درس ها", "url": ""},
            {"label": "فصل ها", "url": f"/courses/myseason/{self.course.id}"},
            {"label": "فایل های من", "url": "/courses/"},
            {"label": "داشبورد", "url": "/"},
        ]

        return {
            'form': form,
            'course': self.course,
            'season': self.season,
            'lessons_data': lessons_data,
            'user_has_all_access': user_has_all_access,
            'user_has_upload_access': user_has_upload_access,
            'breadcrumbs': breadcrumbs
        }

    def get(self, request, *args, **kwargs):
        context = self.get_context_data(request)
        return render(request, self.template_name, context)

    # TODO: Remove this decorator in the final version as it is intended for demo purposes only.
    @restrict_demo_user
    def post(self, request, *args, **kwargs):

        form = LessonForm(request.POST, request.FILES)
        if form.is_valid():
            lesson_content = request.FILES.get('lesson_content').name
            filename, file_extension = os.path.splitext(lesson_content)

            valid_types = ['.mp3', '.mp4', '.avi', '.mkv', '.webm']
            if file_extension not in valid_types:
                messages.error(request, "لطفاً فقط فایل‌های صوتی MP3 یا فایل‌های ویدیویی مجاز آپلود کنید.")
            else:
                lesson_title = filename
                instance = form.save(commit=False)
                instance.course_season = self.season
                instance.title = lesson_title
                instance.order = self.season.fkey_lessons.count() + 1 #Lesson.objects.filter(course_season=season).count() + 1
                instance.save()
                return redirect(reverse(self.mylesson_url, kwargs={'id': self.season.id}))
        else:
            messages.error(request, "لطفاً فرم را به درستی پر کنید.")

        context = self.get_context_data(request)
        return render(request, self.template_name, context)

class CourseUpdateView(LoginRequiredMixin, View):
    login_url = 'login'
    template_name = "courses/user-page/course-update.html"
    mycourse_url = 'mycourse'
    courseupdate_url = 'courseupdate'

    def setup(self, request, *args, **kwargs):
        super().setup(request, *args, **kwargs)
        try:   # گرفتن دوره و دسترسی به آن
            self.course = Course.objects.get(pk=kwargs['id'])
            self.course_access = CourseAccess.objects.get(course=self.course)
        except Course.DoesNotExist:
            raise Http404("دوره انتخابی وجود ندارد")
        except CourseAccess.DoesNotExist:
            raise Http404("دسترسی به دوره وجود ندارد")

    def get_context_data(self, request):

        user_has_all_access = request.user.profileuser.has_permission('All-Access')
        user_has_delete_access = request.user.profileuser.has_permission('Delete')
        user_has_course_access = request.user.profileuser.has_permission('Course-Access')

        form = CreateCourseForm(instance=self.course)
        form_access = CourseAccessForm(instance=self.course_access)

        breadcrumbs = [
            {"label": "ویرایش دوره", "url": ""},
            {"label": "فایل های من", "url": "/courses/"},
            {"label": "داشبورد", "url": "/"},
        ]

        return {
            'form': form,
            'form_access': form_access,
            'course': self.course,
            'user_has_all_access': user_has_all_access,
            'user_has_delete_access': user_has_delete_access,
            'user_has_course_access': user_has_course_access,
            'breadcrumbs': breadcrumbs
        }

    def get(self, request, *args, **kwargs):

        context = self.get_context_data(request)
        return render(request, self.template_name, context)

    def post(self, request, *args, **kwargs):

        # TODO: Remove this decorator in the final version as it is intended for demo purposes only.
        if is_demo_user(request):
            context = self.get_context_data(request)
            return render(request, self.template_name, context)

        action = request.POST.get('action')
        if action == 'OkayDelete':
            if self.course.lesson_count == 0:
                self.course.delete()
                return redirect(reverse(self.mycourse_url))

        elif action == 'OkayUpdate':
            form = CreateCourseForm(request.POST, request.FILES, instance=self.course)
            if form.is_valid() and form.has_changed():
                form.save()
                return redirect(reverse(self.courseupdate_url, kwargs={'id': self.course.id}))
            else:
                messages.error(request, 'لطفاً فرم را به درستی پر کنید.')

        elif action == 'OkayAccess':
            form_access = CourseAccessForm(request.POST, request.FILES, instance=self.course_access)
            if form_access.is_valid() and form_access.has_changed():
                form_access.save()
            else:
                messages.error(request, 'لطفاً فرم را به درستی پر کنید.')

        context = self.get_context_data(request)
        return render(request, self.template_name, context)

class SeasonUpdateView(LoginRequiredMixin, View):
    login_url = 'login'  # مسیر لاگین در صورت نیاز به ورود
    template_name = "courses/user-page/season-update.html"
    myseason_url = 'myseason'
    seasonupdate_url = 'seasonupdate'

    def setup(self, request, *args, **kwargs):
        super().setup(request, *args, **kwargs)
        try:
            self.season = CourseSeason.objects.get(pk=kwargs['id'])
            self.course = self.season.course
        except CourseSeason.DoesNotExist:
            raise Http404("فصل انتخابی وجود ندارد")
        except AttributeError:
            raise Http404("دوره مربوط به فصل وجود ندارد")

    def get_context_data(self, request):

        user_has_all_access = request.user.profileuser.has_permission('All-Access')
        user_has_delete_access = request.user.profileuser.has_permission('Delete')

        form = CourseSeasonForm(instance=self.season)

        breadcrumbs = [
            {"label": "ویرایش فصل", "url": ""},
            {"label": "فصل ها", "url": f"/courses/myseason/{self.course.id}"},
            {"label": "فایل های من", "url": "/courses/"},
            {"label": "داشبورد", "url": "/"},
        ]

        return {
            'form': form,
            'course': self.course,
            'season': self.season,
            'user_has_all_access': user_has_all_access,
            'user_has_delete_access': user_has_delete_access,
            'breadcrumbs': breadcrumbs
        }

    def get(self, request,*args, **kwargs):

        context = self.get_context_data(request)
        return render(request, self.template_name, context)

    def post(self, request, *args, **kwargs):

        action = request.POST.get('action')

        if action == 'OkayDelete':
            if self.season.order != 1:  # جلوگیری از حذف فصل اول
                self.season.delete()
                return redirect(reverse(self.myseason_url, kwargs={'id': self.course.id}))

        elif action == 'OkayUpdate':
            form = CourseSeasonForm(request.POST, request.FILES, instance=self.season)
            if form.is_valid() and form.has_changed():
                form.save()
                return redirect(reverse(self.seasonupdate_url, kwargs={'id': self.season.id}))
            else:
                messages.error(request, 'لطفاً فرم را به درستی پر کنید.')

        context = self.get_context_data(request)
        return render(request, self.template_name, context)

class LessonUpdateView(LoginRequiredMixin, View):
    login_url = 'login'
    template_name = "courses/user-page/lesson-update.html"
    lesson_url = 'mylesson'
    lessonupdate_url = 'lessonupdate'

    def setup(self, request, *args, **kwargs):
        super().setup(request, *args, **kwargs)
        try:
            self.lesson = Lesson.objects.get(pk=kwargs['id'])
            self.season = self.lesson.course_season
            self.course = self.season.course
        except Lesson.DoesNotExist:
            raise Http404("آموزش انتخابی وجود ندارد")
        except AttributeError:
            raise Http404("دوره یا فصل مربوط به درس یافت نشد")

    def get_context_data(self, request):

        user_has_all_access = request.user.profileuser.has_permission('All-Access')
        user_has_delete_access = request.user.profileuser.has_permission('Delete')

        form = LessonForm(instance=self.lesson)

        breadcrumbs = [
            {"label": "ویرایش درس", "url": ""},
            {"label": "درس ها", "url": f"/courses/mylesson/{self.lesson.id}"},
            {"label": "فصل ها", "url": f"/courses/myseason/{self.course.id}"},
            {"label": "فایل های من", "url": "/courses/"},
            {"label": "داشبورد", "url": "/"},
        ]

        return {
            'form': form,
            'lesson': self.lesson,
            'course': self.course,
            'season': self.season,
            'user_has_all_access': user_has_all_access,
            'user_has_delete_access': user_has_delete_access,
            'breadcrumbs': breadcrumbs
        }

    def get(self, request, *args, **kwargs):
        context = self.get_context_data(request)
        return render(request, self.template_name, context)

    # TODO: Remove this decorator in the final version as it is intended for demo purposes only.
    @restrict_demo_user
    def post(self, request, *args, **kwargs):

        action = request.POST.get('action')

        if action == 'OkayDelete':
            try:
                self.lesson.delete()
            except ProtectedError:
                messages.error(request, 'Cannot delete: There are dependent records.')
            except Exception as e:
                messages.error(request, f"پیغام خطا: {e}")

            return redirect(reverse(self.lesson_url, kwargs={'id': self.season.id}))

        elif action == 'OkayUpdate':
            form = LessonForm(request.POST, request.FILES, instance=self.lesson)
            if form.is_valid() and form.has_changed():
                try:
                    form.save()
                except Exception as e:
                    messages.error(request, f"پیغام خطا: {e}")
                return redirect(reverse(self.lessonupdate_url, kwargs={'id': self.lesson.id}))
            else:
                messages.error(request, 'لطفاً فرم را به درستی پر کنید.')

        context = self.get_context_data(request)
        return render(request, self.template_name, context)

class QuoteView(View):
    template_name = "courses/quote.html"

    def get(self, request, *args, **kwargs):

        form_search = SearchForm(request.GET)
        quotes_data = []
        link = request.GET.get('link', '')

        if link == "like":
            quotes_data = Quote.objects.annotate(like_count=Count('likes')).filter(like_count__gt=0).order_by('-like_count')

        if not quotes_data:
            all_ids = list(Quote.objects.values_list('id', flat=True))
            if all_ids:
                random_ids = random.sample(all_ids, min(62, len(all_ids)))
                quotes_data = Quote.objects.filter(id__in=random_ids)

        if form_search.is_valid():
            search_quote = form_search.cleaned_data.get("SearchText")
            if search_quote:
                quotes_data = Quote.objects.filter(
                    Q(quote__icontains=search_quote) | Q(quote_speaker__icontains=search_quote))

        breadcrumbs = [
            {"label": "نقه قول ها", "url": ""},
            {"label": "داشبورد", "url": "/"},
        ]

        context = {
            'form_search': form_search,
            'quotes_data': quotes_data,
            'breadcrumbs': breadcrumbs
        }
        return render(request, self.template_name, context)

class MyCertificateView(LoginRequiredMixin, View):
    login_url = 'login'
    template_name = "courses/user-page/mycertificate.html"

    def get_context_data(self, request):
        try:
            progressed_courses = StudentProgress.objects.filter(
                student=request.user, completed_lessons_count__gt=0
            ).select_related('course')
            courses_data = []

            for progress in progressed_courses:
                course = progress.course
                percent_completed = (progress.completed_lessons_count / course.lesson_count) * 100
                certificate = Certificate.objects.filter(student=request.user, course=course).first()

                courses_data.append({
                    'course': course,
                    'percent_completed': round(percent_completed, 0),
                    'completed_lessons': progress.completed_lessons_count,
                    'certificate_status': certificate.certificate_status if certificate else 'No certificate',
                    'certificate_file': certificate.certificate_file.url if certificate and certificate.certificate_file else None,
                    'certificate_code': certificate.certificate_code if certificate else None,
                })

        except Exception:
            raise Http404("مشکلی در بارگذاری اطلاعات پیش آمده است")

        # Pagination
        courses_data = paginate_queryset(courses_data, request, per_page=6)

        breadcrumbs = [
            {"label": "مدارک من", "url": ""},
            {"label": "داشبورد", "url": "/"},
        ]

        return {
            'courses_data': courses_data,
            'breadcrumbs': breadcrumbs
        }

    def get(self, request):
        context = self.get_context_data(request)
        return render(request, self.template_name, context)


@login_required
def certificate_requestView(request, id):
    course = get_object_or_404(Course, id=id)
    certificate, created = Certificate.objects.get_or_create(student=request.user, course=course,
                                                             defaults={'certificate_status': 'pending'})

    if not StudentProgress.objects.filter(student=request.user, course=course,completed_lessons_count=course.lesson_count).exists():
        raise Http404("شما این دوره را تکمیل نکرده‌اید")

    if not request.user.get_full_name().strip():
        raise Http404("برای دریافت گواهی‌نامه دوره، لطفاً نام و نام خانوادگی خود را در پروفایل وارد کنید")

    # تولید فایل گواهی
    try:
        template_path = certificate.generate_certificate_template()
        doc = Document(template_path)

        # جایگذاری مقادیر در فایل قالب
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if "{{student}}" in run.text:
                    run.text = run.text.replace("{{student}}", request.user.get_full_name())
                if "{{course}}" in run.text:
                    run.text = run.text.replace("{{course}}", course.title)
                if "{{date}}" in run.text:
                    run.text = run.text.replace("{{date}}", certificate.date_of_issue.strftime('%Y-%m-%d'))
                if "{{code}}" in run.text:
                    run.text = run.text.replace("{{code}}", certificate.certificate_code)


        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for run in cell.paragraphs[0].runs:  # دسترسی به run های هر سلول
                        if "{{student}}" in run.text:
                            run.text = run.text.replace("{{student}}", request.user.get_full_name())
                        if "{{course}}" in run.text:
                            run.text = run.text.replace("{{course}}", course.title)
                        if "{{date}}" in run.text:
                            run.text = run.text.replace("{{date}}", certificate.date_of_issue.strftime('%Y-%m-%d'))
                        if "{{code}}" in run.text:
                            run.text = run.text.replace("{{code}}", certificate.certificate_code)

        output_dir = os.path.join(settings.MEDIA_ROOT, 'certificates')
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, f"{request.user.id}_{id}_{request.user.username}_certificate.docx")
        doc.save(output_path)

        # به‌روزرسانی مدل گواهی
        certificate.certificate_file.name = f"certificates/{request.user.id}_{id}_{request.user.username}_certificate.docx"
        certificate.certificate_status = 'issued'
        certificate.save()

        return redirect('mycertificate')

    except FileNotFoundError as e:
        raise Http404(f"خطا: {str(e)}")

@login_required
def certificate_pdfView(request, id):
    course = get_object_or_404(Course, id=id)
    try:
        certificate = Certificate.objects.get(student=request.user, course=course)
    except Certificate.DoesNotExist:
        raise Http404("گواهی نامه وجود ندارد")

    output_dir = os.path.join(settings.MEDIA_ROOT, 'certificates')
    pdf_path = os.path.join(output_dir, f"{request.user.id}_{id}_{request.user.username}_certificate.pdf")

    certificate_file_path = certificate.certificate_file.path  # مسیر واقعی فایل
    if not os.path.exists(pdf_path):
        try:
           convert(certificate_file_path, pdf_path)
        except Exception as e:
            raise Http404(f"Error converting DOCX to PDF: {e}")

    with open(pdf_path, 'rb') as file:
        response = HttpResponse(file.read(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{request.user.id}_{id}_{request.user.username}_certificate.pdf"'
        return response


# ------------------------------------------------------------------------------- Serializers & APIView

# Serializers
class LikeSerializer(serializers.Serializer):
    status = serializers.CharField()
    total_likes = serializers.IntegerField(required=False)

class CompletedSerializer(serializers.Serializer):
    completed = serializers.BooleanField()

class TagSerializer(serializers.Serializer):
    status = serializers.CharField()

# API Views
class LessonLike_APIView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, id):
        lesson = get_object_or_404(Lesson, id=id)

        if request.user in lesson.likes.all():
            lesson.likes.remove(request.user)
            status = "unliked"
        else:
            lesson.likes.add(request.user)
            status = "liked"

        serializer = LikeSerializer({
            'status': status,
            'total_likes': lesson.likes.count(),
        })
        return Response(serializer.data)


class LessonCompleted_APIView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, id):
        lesson = get_object_or_404(Lesson, id=id)

        if request.user in lesson.completed.all():
            lesson.completed.remove(request.user)
            completed = False
        else:
            lesson.completed.add(request.user)
            completed = True

        serializer = CompletedSerializer({'completed': completed})
        return Response(serializer.data)

class CourseLike_APIView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, id):
        course = get_object_or_404(Course, id=id)

        if request.user in course.likes.all():
            course.likes.remove(request.user)
            status = "unliked"
        else:
            course.likes.add(request.user)
            status = "liked"

        serializer = LikeSerializer({
            'status': status,
            'total_likes': course.likes.count(),
        })
        return Response(serializer.data)

class CourseTag_APIView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, id):
        course = get_object_or_404(Course, id=id)

        if request.user in course.tag.all():
            course.tag.remove(request.user)
            status = "untag"
        else:
            course.tag.add(request.user)
            status = "taged"

        serializer = TagSerializer({'status': status})
        return Response(serializer.data)

class QuoteLike_APIView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, id):
        quote = get_object_or_404(Quote, id=id)

        if request.user in quote.likes.all():
            quote.likes.remove(request.user)
            status = "unliked"
        else:
            quote.likes.add(request.user)
            status = "liked"

        serializer = LikeSerializer({
            'status': status,
            'total_likes': quote.likes.count(),
        })
        return Response(serializer.data)




